"""SonettoHere files tools — migrated to SuperAgent plugin.

Read-only file tools (no approval needed):
- pdf_reader: 读取 PDF 文件 (需要 PyPDF2)
- doc_reader: 读取 Word 文档 (需要 python-docx)

Both gracefully report missing dependencies.
"""

from __future__ import annotations

import json
import os
from typing import Any, Dict

from app.tools.base import Tool, ToolContext, ToolResult, ToolSpec


def _ok(data: dict) -> str:
    return json.dumps({"success": True, "data": data}, ensure_ascii=False)


def _err(message: str) -> str:
    return json.dumps({"success": False, "error": message}, ensure_ascii=False)


class PDFReaderTool(Tool):
    spec = ToolSpec(
        name="pdf_reader",
        description=(
            "读取 PDF 文件：元数据、文本提取、关键词搜索、目录、页数。"
            "需要安装 PyPDF2（pip install PyPDF2）。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "operation": {
                    "type": "string",
                    "enum": ["get_metadata", "extract_text", "search_text", "get_toc", "get_page_count"],
                    "description": "操作类型",
                },
                "file_path": {"type": "string", "description": "项目内 PDF 文件相对路径"},
                "start_page": {"type": "integer", "description": "起始页码（从0开始），默认 0"},
                "end_page": {"type": "integer", "description": "结束页码（从0开始），默认最后一页"},
                "query": {"type": "string", "description": "搜索关键词（search_text 操作）"},
                "max_length": {"type": "integer", "description": "文本最大长度，默认 10000"},
            },
            "required": ["operation", "file_path"],
        },
        doc="# pdf_reader\n\n读取 PDF 文件。支持 get_metadata/extract_text/search_text/get_toc/get_page_count。\n需要 PyPDF2：`pip install PyPDF2`",
    )

    async def run(self, args: Dict[str, Any], ctx: ToolContext) -> ToolResult:
        operation = str(args.get("operation") or "")
        file_path = str(args.get("file_path") or "")
        start_page = int(args.get("start_page") or 0)
        end_page = args.get("end_page")
        query = str(args.get("query") or "")
        max_length = int(args.get("max_length") or 10000)

        if not operation or not file_path:
            return ToolResult(False, "", "operation 和 file_path 不能为空")

        from app.tools.workspace import resolve_in_workspace, PathEscapeError
        try:
            p = resolve_in_workspace(ctx.workspace_root, file_path)
        except PathEscapeError as exc:
            return ToolResult(False, "", str(exc))
        if not p.is_file():
            return ToolResult(False, "", f"文件不存在: {file_path}")
        if not str(p).lower().endswith(".pdf"):
            return ToolResult(False, "", f"不是 PDF 文件: {file_path}")

        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                return ToolResult(False, "", "pypdf 未安装，请运行: pip install pypdf")

        try:
            with open(str(p), "rb") as f:
                reader = PdfReader(f)

                if operation == "get_page_count":
                    result = {"page_count": len(reader.pages), "file_path": str(p)}
                    return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

                if operation == "get_metadata":
                    meta = {}
                    if reader.metadata:
                        for key in reader.metadata.keys():
                            val = reader.metadata.get(key, "")
                            if isinstance(val, bytes):
                                val = val.decode("utf-8", errors="ignore")
                            meta[key] = str(val)
                    result = {"metadata": meta, "page_count": len(reader.pages)}
                    return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

                if operation == "get_toc":
                    toc = []
                    def parse_outlines(outlines, level=0):
                        for item in outlines:
                            if isinstance(item, list):
                                parse_outlines(item, level + 1)
                            else:
                                pn = None
                                try:
                                    if hasattr(item, "page") and item.page:
                                        pn = reader.get_page_number(item.page) + 1
                                except Exception:
                                    pass
                                toc.append({"title": getattr(item, "title", str(item)), "level": level, "page_number": pn})
                    try:
                        parse_outlines(reader.outlines)
                    except Exception:
                        pass
                    result = {"toc": toc, "total_pages": len(reader.pages)}
                    return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

                if operation == "extract_text":
                    total = len(reader.pages)
                    end = end_page if end_page is not None and end_page < total else total - 1
                    start = max(0, start_page)
                    if start >= total:
                        return ToolResult(False, "", f"起始页码超出范围: {start}，总页数: {total}")
                    parts = []
                    for pn in range(start, end + 1):
                        pt = reader.pages[pn].extract_text()
                        if pt:
                            parts.append(f"--- 第 {pn + 1} 页 ---\n{pt}")
                    text = "\n\n".join(parts)
                    if len(text) > max_length:
                        text = text[:max_length] + "\n\n[内容已截断]"
                    result = {"text": text, "page_range": [start, end], "total_pages": total, "text_length": len(text)}
                    return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

                if operation == "search_text":
                    if not query:
                        return ToolResult(False, "", "请提供搜索关键词")
                    total = len(reader.pages)
                    q = query.lower()
                    results = []
                    for pn in range(total):
                        pt = reader.pages[pn].extract_text()
                        if not pt:
                            continue
                        if q in pt.lower():
                            matched = [
                                {"line_number": i + 1, "content": line.strip()}
                                for i, line in enumerate(pt.split("\n"))
                                if q in line.lower()
                            ]
                            results.append({"page_number": pn + 1, "matched_lines": matched})
                    result = {"query": query, "total_pages": total, "results": results, "total_matches": sum(len(r["matched_lines"]) for r in results)}
                    return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

                return ToolResult(False, "", f"未知操作: {operation}")
        except Exception as exc:
            return ToolResult(False, "", f"PDF 处理失败: {exc}")


class DocReaderTool(Tool):
    spec = ToolSpec(
        name="doc_reader",
        description=(
            "读取 Word 文档（.docx）：元数据、文本提取、段落、表格、关键词搜索。"
            "需要安装 python-docx（pip install python-docx）。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "operation": {
                    "type": "string",
                    "enum": ["get_metadata", "extract_text", "search_text", "get_paragraphs", "get_tables"],
                    "description": "操作类型",
                },
                "file_path": {"type": "string", "description": "项目内 DOCX 文件相对路径"},
                "start_paragraph": {"type": "integer", "description": "起始段落索引（从0开始）"},
                "end_paragraph": {"type": "integer", "description": "结束段落索引（从0开始）"},
                "query": {"type": "string", "description": "搜索关键词（search_text 操作）"},
                "max_length": {"type": "integer", "description": "文本最大长度，默认 10000"},
            },
            "required": ["operation", "file_path"],
        },
        doc="# doc_reader\n\n读取 Word 文档。支持 get_metadata/extract_text/search_text/get_paragraphs/get_tables。\n需要 python-docx：`pip install python-docx`",
    )

    async def run(self, args: Dict[str, Any], ctx: ToolContext) -> ToolResult:
        operation = str(args.get("operation") or "")
        file_path = str(args.get("file_path") or "")
        start_para = int(args.get("start_paragraph") or 0)
        end_para = args.get("end_paragraph")
        query = str(args.get("query") or "")
        max_length = int(args.get("max_length") or 10000)

        if not operation or not file_path:
            return ToolResult(False, "", "operation 和 file_path 不能为空")

        from app.tools.workspace import resolve_in_workspace, PathEscapeError
        try:
            p = resolve_in_workspace(ctx.workspace_root, file_path)
        except PathEscapeError as exc:
            return ToolResult(False, "", str(exc))
        if not p.is_file():
            return ToolResult(False, "", f"文件不存在: {file_path}")
        ext = os.path.splitext(str(p))[1].lower()
        if ext not in (".doc", ".docx"):
            return ToolResult(False, "", f"不是 Word 文档: {file_path}")
        if ext == ".doc":
            return ToolResult(False, "", ".doc 格式不支持，请先转换为 .docx")

        try:
            import docx
        except ImportError:
            return ToolResult(False, "", "python-docx 未安装，请运行: pip install python-docx")

        try:
            doc = docx.Document(str(p))

            if operation == "get_metadata":
                cp = doc.core_properties
                meta = {
                    "title": cp.title or "", "author": cp.author or "",
                    "created": str(cp.created) if cp.created else "",
                    "modified": str(cp.modified) if cp.modified else "",
                }
                result = {"metadata": meta, "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
                return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

            if operation == "extract_text":
                total = len(doc.paragraphs)
                end = end_para if end_para is not None and end_para < total else total - 1
                if start_para >= total:
                    return ToolResult(False, "", f"起始段落超出范围: {start_para}，总段落数: {total}")
                parts = [doc.paragraphs[i].text for i in range(start_para, end + 1) if doc.paragraphs[i].text.strip()]
                text = "\n\n".join(parts)
                if len(text) > max_length:
                    text = text[:max_length] + "\n\n[内容已截断]"
                result = {"text": text, "paragraph_range": [start_para, end], "total_paragraphs": total, "text_length": len(text)}
                return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

            if operation == "get_paragraphs":
                total = len(doc.paragraphs)
                end = end_para if end_para is not None and end_para < total else total - 1
                if start_para >= total:
                    return ToolResult(False, "", f"起始段落超出范围: {start_para}，总段落数: {total}")
                paras = [
                    {"index": i, "text": doc.paragraphs[i].text, "style": doc.paragraphs[i].style.name if doc.paragraphs[i].style else None}
                    for i in range(start_para, end + 1)
                ]
                result = {"paragraphs": paras, "paragraph_range": [start_para, end], "total_paragraphs": total}
                return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

            if operation == "get_tables":
                tables = []
                for idx, table in enumerate(doc.tables):
                    data = [[cell.text for cell in row.cells] for row in table.rows]
                    tables.append({"index": idx, "rows": len(table.rows), "columns": len(table.columns), "data": data})
                result = {"tables": tables, "total_tables": len(tables)}
                return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

            if operation == "search_text":
                if not query:
                    return ToolResult(False, "", "请提供搜索关键词")
                q = query.lower()
                results = []
                for idx, para in enumerate(doc.paragraphs):
                    if para.text and q in para.text.lower():
                        results.append({"paragraph_index": idx, "content": para.text.strip()})
                result = {"query": query, "total_paragraphs": len(doc.paragraphs), "results": results, "total_matches": len(results)}
                return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

            return ToolResult(False, "", f"未知操作: {operation}")
        except Exception as exc:
            return ToolResult(False, "", f"Word 文档处理失败: {exc}")


def register(registry):
    registry.register(PDFReaderTool())
    registry.register(DocReaderTool())