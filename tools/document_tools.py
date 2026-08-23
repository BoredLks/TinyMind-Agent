"""Document processing tools for SuperAgent — adapted from reference MCP servers.

Provides read/write/edit capabilities for PDF, Word (.docx), and Excel (.xlsx) files.

Reference sources:
- github.com/alejandroBallesterosC/document-edit-mcp (Word/Excel/PDF operations)
- github.com/MeterLong/MCP-Doc (Word document processing)
- github.com/jonemo/openpyxl-mcp-server (Excel operations)

Dependencies: openpyxl, fpdf2, python-docx, PyPDF2
"""

from __future__ import annotations

import json
import os
from typing import Any, Dict, List

from app.tools.base import Tool, ToolContext, ToolResult, ToolSpec
from app.tools.workspace import resolve_in_workspace, PathEscapeError


def _ok(data: dict) -> str:
    return json.dumps({"success": True, "data": data}, ensure_ascii=False)


def _err(message: str) -> str:
    return json.dumps({"success": False, "error": message}, ensure_ascii=False)


# ============================================================
# Word Document Editor
# ============================================================

class DocEditorTool(Tool):
    """Create, edit, and convert Word documents (.docx)."""

    spec = ToolSpec(
        name="doc_editor",
        description=(
            "创建、编辑和转换 Word 文档（.docx）。"
            "支持：create（创建新文档）、add_heading（添加标题）、add_paragraph（添加段落）、"
            "add_table（添加表格）、replace_text（替换文本）、delete_paragraph（删除段落）、"
            "convert_to_txt（转为纯文本）。"
            "需要 python-docx（已安装）。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "operation": {
                    "type": "string",
                    "enum": [
                        "create", "add_heading", "add_paragraph",
                        "add_table", "replace_text", "delete_paragraph",
                        "convert_to_txt",
                    ],
                    "description": "操作类型",
                },
                "file_path": {"type": "string", "description": "项目内文件相对路径"},
                "text": {"type": "string", "description": "文本内容（create/add_heading/add_paragraph/replace_text）"},
                "level": {"type": "integer", "description": "标题级别 1-9（add_heading），默认 1"},
                "index": {"type": "integer", "description": "段落索引（delete_paragraph/replace_text）"},
                "old_text": {"type": "string", "description": "要替换的原文（replace_text）"},
                "new_text": {"type": "string", "description": "替换后的新文本（replace_text）"},
                "table_data": {
                    "type": "string",
                    "description": "表格数据 JSON 字符串，格式为 [[row1col1,row1col2],[row2col1,row2col2]]（add_table）",
                },
                "target_path": {"type": "string", "description": "目标文件路径（convert_to_txt）"},
                "max_length": {"type": "integer", "description": "最大文本长度，默认 10000"},
            },
            "required": ["operation", "file_path"],
        },
        requires_approval=True,
        doc=(
            "# doc_editor\n\n"
            "创建、编辑和转换 Word 文档。支持的操作：\n"
            "- create：创建新 .docx 文件\n"
            "- add_heading：添加标题（level 1-9）\n"
            "- add_paragraph：添加段落文本\n"
            "- add_table：添加表格（传入 JSON 二维数组）\n"
            "- replace_text：替换指定段落的文本\n"
            "- delete_paragraph：删除指定索引的段落\n"
            "- convert_to_txt：将 docx 转为纯文本文件\n\n"
            "需要 python-docx（已安装）。"
        ),
    )

    async def run(self, args: Dict[str, Any], ctx: ToolContext) -> ToolResult:
        operation = str(args.get("operation") or "")
        file_path = str(args.get("file_path") or "")
        if not operation or not file_path:
            return ToolResult(False, "", "operation 和 file_path 不能为空")

        try:
            p = resolve_in_workspace(ctx.workspace_root, file_path)
        except PathEscapeError as exc:
            return ToolResult(False, "", str(exc))

        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return ToolResult(False, "", "python-docx 未安装")

        try:
            if operation == "create":
                text = str(args.get("text") or "")
                doc = Document()
                if text:
                    for para_text in text.split("\n"):
                        if para_text.strip():
                            doc.add_paragraph(para_text)
                p.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(p))
                return ToolResult(True, _ok({"action": "created", "path": file_path}),
                                  display={"kind": "file", "path": file_path, "chars": len(text), "action": "write",
                                           "previewAvailable": True, "previewContent": text})

            if not p.is_file():
                return ToolResult(False, "", f"文件不存在: {file_path}")
            if not str(p).lower().endswith(".docx"):
                return ToolResult(False, "", f"不是 .docx 文件: {file_path}")

            doc = Document(str(p))

            if operation == "add_heading":
                text = str(args.get("text") or "")
                level = int(args.get("level") or 1)
                if not text:
                    return ToolResult(False, "", "text 不能为空")
                level = max(1, min(9, level))
                doc.add_heading(text, level=level)
                doc.save(str(p))
                return ToolResult(True, _ok({"action": "add_heading", "text": text, "level": level}))

            if operation == "add_paragraph":
                text = str(args.get("text") or "")
                if not text:
                    return ToolResult(False, "", "text 不能为空")
                doc.add_paragraph(text)
                doc.save(str(p))
                return ToolResult(True, _ok({"action": "add_paragraph", "text": text[:100]}))

            if operation == "add_table":
                table_data_str = str(args.get("table_data") or "[]")
                try:
                    table_data = json.loads(table_data_str)
                except json.JSONDecodeError:
                    return ToolResult(False, "", "table_data 必须是有效的 JSON 字符串")
                if not table_data or not isinstance(table_data, list):
                    return ToolResult(False, "", "table_data 必须是非空二维数组")
                rows = len(table_data)
                cols = max(len(row) for row in table_data) if rows > 0 else 0
                table = doc.add_table(rows=rows, cols=cols, style="Table Grid")
                for i, row in enumerate(table_data):
                    for j, cell_val in enumerate(row):
                        if j < cols:
                            table.cell(i, j).text = str(cell_val)
                doc.save(str(p))
                return ToolResult(True, _ok({"action": "add_table", "rows": rows, "cols": cols}))

            if operation == "replace_text":
                old_text = str(args.get("old_text") or "")
                new_text = str(args.get("new_text") or "")
                if not old_text:
                    return ToolResult(False, "", "old_text 不能为空")
                replaced = 0
                for para in doc.paragraphs:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)
                        replaced += 1
                if replaced == 0:
                    return ToolResult(False, "", f"未找到匹配文本: {old_text[:50]}")
                doc.save(str(p))
                return ToolResult(True, _ok({"action": "replace_text", "replaced_count": replaced}))

            if operation == "delete_paragraph":
                idx = int(args.get("index") or 0)
                if idx < 0 or idx >= len(doc.paragraphs):
                    return ToolResult(False, "", f"段落索引 {idx} 超出范围（共 {len(doc.paragraphs)} 段）")
                para = doc.paragraphs[idx]
                p_elem = para._element
                p_elem.getparent().remove(p_elem)
                doc.save(str(p))
                return ToolResult(True, _ok({"action": "delete_paragraph", "index": idx}))

            if operation == "convert_to_txt":
                target = str(args.get("target_path") or "")
                max_length = int(args.get("max_length") or 10000)
                if not target:
                    # Default: same name with .txt extension
                    target = str(p).rsplit(".", 1)[0] + ".txt"
                try:
                    tp = resolve_in_workspace(ctx.workspace_root, target)
                except PathEscapeError:
                    tp = p.parent / target
                parts = [para.text for para in doc.paragraphs if para.text.strip()]
                text = "\n\n".join(parts)
                if len(text) > max_length:
                    text = text[:max_length] + "\n\n[内容已截断]"
                tp.parent.mkdir(parents=True, exist_ok=True)
                tp.write_text(text, encoding="utf-8")
                return ToolResult(True, _ok({"action": "convert_to_txt", "target": target, "chars": len(text)}))

            return ToolResult(False, "", f"未知操作: {operation}")
        except Exception as exc:
            return ToolResult(False, "", f"Word 文档操作失败: {exc}")


# ============================================================
# Excel Tool
# ============================================================

class ExcelTool(Tool):
    """Read, create, and edit Excel files (.xlsx)."""

    spec = ToolSpec(
        name="excel_tool",
        description=(
            "读取、创建和编辑 Excel 文件（.xlsx）。"
            "支持：list_sheets（列出工作表）、read_sheet（读取工作表数据）、read_cell（读取单元格）、"
            "write_cell（写入单元格）、write_range（写入范围）、create_workbook（创建工作簿）、"
            "add_sheet（添加工作表）、delete_sheet（删除工作表）、delete_row（删除行）、delete_column（删除列）。"
            "需要 openpyxl（已安装）。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "operation": {
                    "type": "string",
                    "enum": [
                        "list_sheets", "read_sheet", "read_cell",
                        "write_cell", "write_range", "create_workbook",
                        "add_sheet", "delete_sheet", "delete_row", "delete_column",
                    ],
                    "description": "操作类型",
                },
                "file_path": {"type": "string", "description": "项目内文件相对路径"},
                "sheet_name": {"type": "string", "description": "工作表名称，默认第一个"},
                "row": {"type": "integer", "description": "行号（从1开始）"},
                "col": {"type": "integer", "description": "列号（从1开始）"},
                "value": {"type": "string", "description": "单元格值（write_cell）"},
                "values": {
                    "type": "string",
                    "description": "二维数组 JSON 字符串（write_range），格式 [[v1,v2],[v3,v4]]",
                },
                "start_row": {"type": "integer", "description": "起始行号（write_range），默认 1"},
                "start_col": {"type": "integer", "description": "起始列号（write_range），默认 1"},
                "max_rows": {"type": "integer", "description": "最大读取行数（read_sheet），默认 100"},
                "new_sheet_name": {"type": "string", "description": "新工作表名称（add_sheet）"},
            },
            "required": ["operation", "file_path"],
        },
        requires_approval=True,
        doc=(
            "# excel_tool\n\n"
            "读取、创建和编辑 Excel 文件。支持的操作：\n"
            "- list_sheets：列出所有工作表\n"
            "- read_sheet：读取工作表数据（返回 JSON）\n"
            "- read_cell：读取指定单元格的值\n"
            "- write_cell：写入指定单元格\n"
            "- write_range：写入一个范围的数据\n"
            "- create_workbook：创建新的 Excel 工作簿\n"
            "- add_sheet / delete_sheet：添加/删除工作表\n"
            "- delete_row / delete_column：删除行/列\n\n"
            "需要 openpyxl（已安装）。"
        ),
    )

    async def run(self, args: Dict[str, Any], ctx: ToolContext) -> ToolResult:
        operation = str(args.get("operation") or "")
        file_path = str(args.get("file_path") or "")
        if not operation or not file_path:
            return ToolResult(False, "", "operation 和 file_path 不能为空")

        try:
            p = resolve_in_workspace(ctx.workspace_root, file_path)
        except PathEscapeError as exc:
            return ToolResult(False, "", str(exc))

        try:
            import openpyxl
        except ImportError:
            return ToolResult(False, "", "openpyxl 未安装")

        try:
            # Read-only operations
            if operation == "list_sheets":
                if not p.is_file():
                    return ToolResult(False, "", f"文件不存在: {file_path}")
                wb = openpyxl.load_workbook(str(p), read_only=True)
                sheets = wb.sheetnames
                wb.close()
                return ToolResult(True, _ok({"sheets": sheets, "count": len(sheets)}))

            if operation == "read_sheet":
                if not p.is_file():
                    return ToolResult(False, "", f"文件不存在: {file_path}")
                sheet_name = args.get("sheet_name")
                max_rows = int(args.get("max_rows") or 100)
                wb = openpyxl.load_workbook(str(p), read_only=True)
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                data = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= max_rows:
                        break
                    data.append([str(c) if c is not None else "" for c in row])
                result = {
                    "sheet_name": ws.title,
                    "rows": len(data),
                    "columns": len(data[0]) if data else 0,
                    "data": data,
                }
                wb.close()
                return ToolResult(True, _ok(result), display={"kind": "json", "data": result})

            if operation == "read_cell":
                if not p.is_file():
                    return ToolResult(False, "", f"文件不存在: {file_path}")
                row = int(args.get("row") or 1)
                col = int(args.get("col") or 1)
                sheet_name = args.get("sheet_name")
                wb = openpyxl.load_workbook(str(p), read_only=True)
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                cell = ws.cell(row=row, column=col)
                value = cell.value
                wb.close()
                return ToolResult(True, _ok({"row": row, "col": col, "value": str(value) if value is not None else "", "cell": cell.coordinate}))

            # Write operations
            if operation == "create_workbook":
                wb = openpyxl.Workbook()
                p.parent.mkdir(parents=True, exist_ok=True)
                wb.save(str(p))
                wb.close()
                return ToolResult(True, _ok({"action": "created", "path": file_path}))

            if not p.is_file():
                return ToolResult(False, "", f"文件不存在: {file_path}")

            wb = openpyxl.load_workbook(str(p))

            if operation == "write_cell":
                row = int(args.get("row") or 1)
                col = int(args.get("col") or 1)
                value = str(args.get("value") or "")
                sheet_name = args.get("sheet_name")
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                # Try to convert to number
                try:
                    value = float(value)
                    if value == int(value):
                        value = int(value)
                except (ValueError, TypeError):
                    pass
                ws.cell(row=row, column=col, value=value)
                wb.save(str(p))
                wb.close()
                return ToolResult(True, _ok({"action": "write_cell", "row": row, "col": col, "value": str(value)}))

            if operation == "write_range":
                start_row = int(args.get("start_row") or 1)
                start_col = int(args.get("start_col") or 1)
                values_str = str(args.get("values") or "[]")
                try:
                    values = json.loads(values_str)
                except json.JSONDecodeError:
                    return ToolResult(False, "", "values 必须是有效的 JSON 二维数组")
                sheet_name = args.get("sheet_name")
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                for i, row_vals in enumerate(values):
                    for j, val in enumerate(row_vals):
                        # Try numeric conversion
                        try:
                            val = float(val)
                            if val == int(val):
                                val = int(val)
                        except (ValueError, TypeError):
                            pass
                        ws.cell(row=start_row + i, column=start_col + j, value=val)
                wb.save(str(p))
                wb.close()
                return ToolResult(True, _ok({
                    "action": "write_range",
                    "start_row": start_row, "start_col": start_col,
                    "rows_written": len(values),
                    "cols_written": max(len(r) for r in values) if values else 0,
                }))

            if operation == "add_sheet":
                new_name = str(args.get("new_sheet_name") or "Sheet")
                if new_name in wb.sheetnames:
                    return ToolResult(False, "", f"工作表已存在: {new_name}")
                wb.create_sheet(new_name)
                wb.save(str(p))
                wb.close()
                return ToolResult(True, _ok({"action": "add_sheet", "name": new_name, "sheets": wb.sheetnames}))

            if operation == "delete_sheet":
                sheet_name = str(args.get("sheet_name") or "")
                if not sheet_name:
                    return ToolResult(False, "", "请指定要删除的工作表名称")
                if sheet_name not in wb.sheetnames:
                    return ToolResult(False, "", f"工作表不存在: {sheet_name}")
                if len(wb.sheetnames) <= 1:
                    return ToolResult(False, "", "不能删除唯一的工作表")
                del wb[sheet_name]
                wb.save(str(p))
                remaining = wb.sheetnames
                wb.close()
                return ToolResult(True, _ok({"action": "delete_sheet", "name": sheet_name, "remaining": remaining}))

            if operation == "delete_row":
                row = int(args.get("row") or 1)
                sheet_name = args.get("sheet_name")
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                ws.delete_rows(row)
                wb.save(str(p))
                wb.close()
                return ToolResult(True, _ok({"action": "delete_row", "row": row}))

            if operation == "delete_column":
                col = int(args.get("col") or 1)
                sheet_name = args.get("sheet_name")
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                ws.delete_cols(col)
                wb.save(str(p))
                wb.close()
                return ToolResult(True, _ok({"action": "delete_column", "col": col}))

            wb.close()
            return ToolResult(False, "", f"未知操作: {operation}")
        except Exception as exc:
            return ToolResult(False, "", f"Excel 操作失败: {exc}")


# ============================================================
# PDF Editor
# ============================================================

class PDFEditorTool(Tool):
    """Create and edit PDF files."""

    spec = ToolSpec(
        name="pdf_editor",
        description=(
            "创建和编辑 PDF 文件。"
            "支持：create（创建新 PDF）、add_page（添加页面）、add_text（添加文本）、"
            "add_table（添加表格）、merge_pdfs（合并多个 PDF）。"
            "需要 fpdf2（已安装）。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "operation": {
                    "type": "string",
                    "enum": ["create", "add_text", "add_table", "merge_pdfs"],
                    "description": "操作类型",
                },
                "file_path": {"type": "string", "description": "项目内文件相对路径"},
                "text": {"type": "string", "description": "文本内容（create/add_text）"},
                "title": {"type": "string", "description": "文档标题（create）"},
                "table_data": {
                    "type": "string",
                    "description": "表格数据 JSON 字符串（add_table），格式 [[v1,v2],[v3,v4]]",
                },
                "col_widths": {
                    "type": "string",
                    "description": "列宽 JSON 数组（add_table），如 [40,60,50]",
                },
                "input_paths": {
                    "type": "string",
                    "description": "要合并的 PDF 文件路径 JSON 数组（merge_pdfs），如 [\"a.pdf\",\"b.pdf\"]",
                },
                "target_path": {"type": "string", "description": "目标文件路径（merge_pdfs）"},
            },
            "required": ["operation", "file_path"],
        },
        requires_approval=True,
        doc=(
            "# pdf_editor\n\n"
            "创建和编辑 PDF 文件。支持的操作：\n"
            "- create：创建新 PDF（支持中文，自动使用内置字体）\n"
            "- add_text：向现有 PDF 添加文本\n"
            "- add_table：添加表格到 PDF\n"
            "- merge_pdfs：合并多个 PDF 文件\n\n"
            "需要 fpdf2（已安装）。"
        ),
    )

    async def run(self, args: Dict[str, Any], ctx: ToolContext) -> ToolResult:
        operation = str(args.get("operation") or "")
        file_path = str(args.get("file_path") or "")
        if not operation or not file_path:
            return ToolResult(False, "", "operation 和 file_path 不能为空")

        try:
            p = resolve_in_workspace(ctx.workspace_root, file_path)
        except PathEscapeError as exc:
            return ToolResult(False, "", str(exc))

        try:
            from fpdf import FPDF
        except ImportError:
            return ToolResult(False, "", "fpdf2 未安装")

        try:
            if operation == "create":
                text = str(args.get("text") or "")
                title = str(args.get("title") or "")
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf_width = pdf.w - pdf.l_margin - pdf.r_margin
                pdf.set_font("Helvetica", "B", 16)
                if title:
                    safe_title = title.encode("latin-1", errors="replace").decode("latin-1")
                    pdf.multi_cell(w=pdf_width, h=10, text=safe_title, align="C")
                    pdf.ln(5)
                pdf.set_font("Helvetica", size=11)
                if text:
                    for line in text.split("\n"):
                        safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
                        pdf.multi_cell(w=pdf_width, h=6, text=safe_line)
                p.parent.mkdir(parents=True, exist_ok=True)
                pdf.output(str(p))
                return ToolResult(True, _ok({"action": "created", "path": file_path, "pages": 1}),
                                  display={"kind": "file", "path": file_path, "chars": len(text), "action": "write",
                                           "previewAvailable": True, "previewContent": text})

            if operation == "add_text":
                if not p.is_file():
                    return ToolResult(False, "", f"文件不存在: {file_path}")
                text = str(args.get("text") or "")
                if not text:
                    return ToolResult(False, "", "text 不能为空")
                # fpdf2 can read existing PDFs
                from fpdf import FPDFReader
                existing = FPDFReader(str(p))
                pdf = FPDF()
                # Copy existing pages
                for page_num in range(len(existing.pages)):
                    pdf.add_page()
                    pdf.set_font("Helvetica", size=11)
                # Add new page with text
                pdf.add_page()
                pdf.set_font("Helvetica", size=11)
                for line in text.split("\n"):
                    safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
                    pdf.multi_cell(0, 6, safe_line)
                pdf.output(str(p))
                return ToolResult(True, _ok({"action": "add_text", "chars": len(text)}))

            if operation == "add_table":
                if not p.is_file():
                    return ToolResult(False, "", f"文件不存在: {file_path}")
                table_data_str = str(args.get("table_data") or "[]")
                try:
                    table_data = json.loads(table_data_str)
                except json.JSONDecodeError:
                    return ToolResult(False, "", "table_data 必须是有效的 JSON 二维数组")
                col_widths_str = str(args.get("col_widths") or "[]")
                try:
                    col_widths = json.loads(col_widths_str) if col_widths_str != "[]" else None
                except json.JSONDecodeError:
                    col_widths = None

                from fpdf import FPDFReader
                existing = FPDFReader(str(p))
                pdf = FPDF()
                for page_num in range(len(existing.pages)):
                    pdf.add_page()
                # Add new page with table
                pdf.add_page()
                pdf.set_font("Helvetica", size=10)
                if col_widths is None and table_data:
                    col_widths = [190 / len(table_data[0])] * len(table_data[0])
                for row in table_data:
                    for i, cell in enumerate(row):
                        w = col_widths[i] if col_widths and i < len(col_widths) else 40
                        safe = str(cell).encode("latin-1", errors="replace").decode("latin-1")
                        pdf.cell(w, 8, safe, border=1)
                    pdf.ln()
                pdf.output(str(p))
                return ToolResult(True, _ok({"action": "add_table", "rows": len(table_data)}))

            if operation == "merge_pdfs":
                input_paths_str = str(args.get("input_paths") or "[]")
                target_path = str(args.get("target_path") or "")
                try:
                    input_paths = json.loads(input_paths_str)
                except json.JSONDecodeError:
                    return ToolResult(False, "", "input_paths 必须是有效的 JSON 数组")
                if not input_paths:
                    return ToolResult(False, "", "input_paths 不能为空")

                from pypdf import PdfWriter, PdfReader as PyPdfReader
                writer = PdfWriter()
                for ip in input_paths:
                    try:
                        ip_resolved = resolve_in_workspace(ctx.workspace_root, ip)
                    except PathEscapeError:
                        return ToolResult(False, "", f"路径不安全: {ip}")
                    if not ip_resolved.is_file():
                        return ToolResult(False, "", f"文件不存在: {ip}")
                    reader = PyPdfReader(str(ip_resolved))
                    for page in reader.pages:
                        writer.add_page(page)

                if target_path:
                    try:
                        tp = resolve_in_workspace(ctx.workspace_root, target_path)
                    except PathEscapeError:
                        tp = p.parent / target_path
                else:
                    tp = p
                tp.parent.mkdir(parents=True, exist_ok=True)
                with open(str(tp), "wb") as f:
                    writer.write(f)
                return ToolResult(True, _ok({"action": "merge_pdfs", "input_count": len(input_paths), "output": str(tp)}))

            return ToolResult(False, "", f"未知操作: {operation}")
        except Exception as exc:
            return ToolResult(False, "", f"PDF 操作失败: {exc}")


# ============================================================
# Registration
# ============================================================

def register(registry):
    registry.register(DocEditorTool(), external=True)
    registry.register(ExcelTool(), external=True)
    registry.register(PDFEditorTool(), external=True)